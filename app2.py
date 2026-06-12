import streamlit as st
import pandas as pd
import numpy as np
import altair as alt
import re
import os
from io import BytesIO
from datetime import datetime
from docx import Document
from openpyxl import load_workbook


def euro(val):
    return f"EUR {val:,.0f}".replace(",", ".")


def format_number_it(val, decimals=2):
    if pd.isna(val):
        return ""
    s = f"{val:,.{decimals}f}"
    return s.replace(",", "X").replace(".", ",").replace("X", ".")


def parse_number_it(val):
    if pd.isna(val):
        return np.nan
    s = str(val).strip()
    if not s:
        return np.nan
    s = s.replace("\u00a0", " ").replace("EUR", "").replace("eur", "").replace(" ", "").replace("'", "")
    s = s.replace("€", "")
    neg_parentheses = s.startswith("(") and s.endswith(")")
    if neg_parentheses:
        s = s[1:-1].strip()

    has_comma = "," in s
    has_dot = "." in s
    if has_comma and has_dot:
        if s.rfind(",") > s.rfind("."):
            s = s.replace(".", "").replace(",", ".")
        else:
            s = s.replace(",", "")
    elif has_comma:
        s = s.replace(".", "").replace(",", ".")
    elif has_dot:
        parts = s.split(".")
        if len(parts) > 2 and all(p.isdigit() for p in parts):
            s = "".join(parts)

    if not re.match(r"^-?\d+(\.\d+)?$", s):
        return np.nan
    n = float(s)
    return -n if neg_parentheses else n


def normalize_headers(headers):
    normalized = []
    for idx, header in enumerate(headers):
        name = str(header).strip() if header is not None else ""
        normalized.append(name or f"Colonna_{idx + 1}")
    return normalized


def with_display_index(df):
    displayed_df = df.copy()
    displayed_df.index = pd.RangeIndex(start=1, stop=len(displayed_df) + 1, step=1)
    return displayed_df


def format_display_df(df, numeric_formats):
    displayed_df = with_display_index(df)
    for col, decimals in numeric_formats.items():
        if col in displayed_df.columns:
            displayed_df[col] = displayed_df[col].apply(lambda x, d=decimals: format_number_it(x, d))
    return displayed_df


def detect_csv_separator(file_bytes):
    for sep in [",", ";", "\t"]:
        try:
            preview_df = pd.read_csv(BytesIO(file_bytes), sep=sep, nrows=5)
            if len(preview_df.columns) > 1:
                return sep
        except Exception:
            continue
    return None


def render_universe_stratification(df, valore_col):
    if df.empty:
        return

    strata = [
        ("Top 1% (0%-1%)", 0.01),
        (">1%-5%", 0.05),
        (">5%-10%", 0.10),
        (">10%-25%", 0.25),
        (">25%-50%", 0.50),
        (">50%-100%", 1.00),
    ]
    sorted_values = df[[valore_col]].sort_values(by=valore_col, ascending=False).reset_index(drop=True)
    sorted_values["rank_pct"] = (np.arange(len(sorted_values)) + 1) / len(sorted_values)
    previous_limit = 0.0
    rows = []
    for label, limit in strata:
        mask = (sorted_values["rank_pct"] > previous_limit) & (sorted_values["rank_pct"] <= limit)
        stratum_values = sorted_values.loc[mask, valore_col]
        rows.append(
            {
                "Fascia": label,
                "Da rank %": previous_limit * 100,
                "A rank %": limit * 100,
                "Numero items": int(stratum_values.count()),
                "Valore": float(stratum_values.sum()),
            }
        )
        previous_limit = limit

    strat_df = pd.DataFrame(rows)
    strat_df = strat_df[strat_df["Numero items"] > 0].reset_index(drop=True)
    if strat_df.empty:
        return

    st.subheader("Stratificazione universo")
    st.write("Valore per singolo item e valore progressivo, ordinati dal valore piu elevato")
    chart_df = sorted_values[[valore_col]].copy()
    chart_df["Numero item"] = np.arange(1, len(chart_df) + 1)
    chart_df = chart_df.rename(columns={valore_col: "Valore"})
    chart_df["Valore progressivo"] = chart_df["Valore"].cumsum()
    totale_grafico = float(chart_df["Valore"].sum())
    chart_df["Progressivo %"] = np.where(
        totale_grafico != 0,
        chart_df["Valore progressivo"] / totale_grafico * 100,
        0.0,
    )

    bars = (
        alt.Chart(chart_df)
        .mark_bar(color="#2f5597", opacity=0.75)
        .encode(
            x=alt.X("Numero item:Q", title="Numero item"),
            y=alt.Y("Valore:Q", title="Valore singolo item"),
            tooltip=[
                alt.Tooltip("Numero item:Q", title="Numero item", format=",.0f"),
                alt.Tooltip("Valore:Q", title="Valore", format=",.2f"),
                alt.Tooltip("Valore progressivo:Q", title="Valore progressivo", format=",.2f"),
                alt.Tooltip("Progressivo %:Q", title="Progressivo %", format=".2f"),
            ],
        )
    )
    line = (
        alt.Chart(chart_df)
        .mark_line(color="#c55a11", strokeWidth=2.5)
        .encode(
            x=alt.X("Numero item:Q", title="Numero item"),
            y=alt.Y("Valore progressivo:Q", title="Valore progressivo"),
            tooltip=[
                alt.Tooltip("Numero item:Q", title="Numero item", format=",.0f"),
                alt.Tooltip("Valore progressivo:Q", title="Valore progressivo", format=",.2f"),
                alt.Tooltip("Progressivo %:Q", title="Progressivo %", format=".2f"),
            ],
        )
    )
    st.altair_chart((bars + line).resolve_scale(y="independent"), use_container_width=True)
    st.caption(
        "Ogni barra rappresenta un item dell'universo; la linea mostra il valore progressivo cumulato."
    )

    st.dataframe(
        format_display_df(
            strat_df,
            {"Da rank %": 0, "A rank %": 0, "Numero items": 0, "Valore": 0},
        ),
        use_container_width=True,
    )


@st.cache_data(show_spinner=False)
def get_excel_sheet_names(file_bytes):
    wb = load_workbook(BytesIO(file_bytes), read_only=True, data_only=True)
    try:
        return wb.sheetnames
    finally:
        wb.close()


@st.cache_data(show_spinner=False)
def get_file_columns(file_bytes, ext, sheet_name=None):
    if ext == ".xlsx":
        wb = load_workbook(BytesIO(file_bytes), read_only=True, data_only=True)
        try:
            ws = wb[sheet_name] if sheet_name in wb.sheetnames else wb.worksheets[0]
            header_row = next(ws.iter_rows(min_row=1, max_row=1, values_only=True), None)
        finally:
            wb.close()
        if not header_row:
            return []
        return normalize_headers(header_row)

    if ext == ".csv":
        sep = detect_csv_separator(file_bytes)
        if sep is None:
            return []
        csv_df = pd.read_csv(BytesIO(file_bytes), sep=sep, nrows=0)
        return normalize_headers(csv_df.columns.tolist())

    return []


@st.cache_data(show_spinner=False)
def load_selected_columns(file_bytes, ext, selected_columns, sheet_name=None):
    if ext == ".xlsx":
        wb = load_workbook(BytesIO(file_bytes), read_only=True, data_only=True)
        try:
            ws = wb[sheet_name] if sheet_name in wb.sheetnames else wb.worksheets[0]
            header_row = next(ws.iter_rows(min_row=1, max_row=1, values_only=True), None)
            if not header_row:
                return pd.DataFrame(columns=list(selected_columns))

            headers = normalize_headers(header_row)
            col_idx = [headers.index(col) for col in selected_columns]
            rows = []
            for row in ws.iter_rows(min_row=2, values_only=True):
                rows.append(
                    {
                        selected_columns[pos]: row[idx] if idx < len(row) else None
                        for pos, idx in enumerate(col_idx)
                    }
                )
        finally:
            wb.close()
        return pd.DataFrame(rows, columns=list(selected_columns))

    if ext == ".csv":
        sep = detect_csv_separator(file_bytes)
        if sep is None:
            return pd.DataFrame(columns=list(selected_columns))
        return pd.read_csv(BytesIO(file_bytes), sep=sep, usecols=list(selected_columns))

    return pd.DataFrame(columns=list(selected_columns))


st.set_page_config(page_title="Audit Sampling - Key Items & Items", layout="wide")
st.title("Selezione campione - Key Items e Items")
st.caption("Versione con riepilogo Selezioni teoriche / Items unici effettivi")

with st.sidebar.expander("Metodologia e funzionamento del programma"):
    st.markdown(
        """
**DESCRIZIONE DEL PROGRAMMA - AUDIT SAMPLING**

**1. Finalita del programma**
Il programma Audit Sampling - Key Items & Items e uno strumento di supporto alle attivita di revisione contabile, sviluppato per la selezione strutturata dei campioni secondo i principi degli ISA (International Standards on Auditing).

Consente di:
- identificare automaticamente i Key Items
- selezionare un campione di elementi residui
- applicare criteri coerenti con la materialita e il livello di rischio
- garantire riproducibilita e tracciabilita delle scelte

**2. Riferimenti normativi (ISA)**
Il programma e coerente con:
- ISA 300 - Pianificazione della revisione
- ISA 315 - Identificazione e valutazione dei rischi
- ISA 330 - Risposte ai rischi identificati
- ISA 500 - Elementi probativi
- ISA 530 - Campionamento di revisione

**3. Logica generale di funzionamento**
Il processo di selezione si articola in tre fasi:
- Individuazione dei Key Items: elementi di maggiore rilevanza monetaria rispetto al totale.
- Determinazione del residuo: popolazione residua al netto dei Key Items.
- Selezione del campione: mediante uno dei metodi disponibili (MUS, Intervallo, Casuale).

**4. Metodo MUS (Monetary Unit Sampling)**
Il metodo MUS:
- assegna maggiore probabilita agli importi piu elevati
- e coerente con test di esistenza e accuratezza
- e particolarmente adatto per crediti, ricavi, immobilizzazioni
- utilizza un seed fisso (42) per garantire riproducibilita

Il programma calcola automaticamente:
- intervallo di campionamento
- starting point
- numero di elementi da testare

**5. Parametri di selezione**
L'utente definisce:
- Soglia Key Items (%)
- Materialita (EUR)
- Confidence level (%)
- Metodo di selezione

Tutti i parametri sono visibili e tracciabili.

**6. Output del programma**
Il software produce:
- Riepilogo generale
- Elenco Key Items
- Campione selezionato
- Export Excel
- Export Word

Gli importi sono:
- espressi in euro
- senza decimali
- coerenti con la documentazione di revisione

**7. Riproducibilita e controlli**
- Seed fisso per evitare variazioni casuali
- Nessuna modifica ai dati originali
- Risultati coerenti a parita di input
- Processo pienamente tracciabile

**8. Finalita di revisione**
Il programma supporta:
- la pianificazione delle verifiche
- la definizione del campione
- la documentazione delle scelte
- la difendibilita del lavoro svolto
- l'allineamento agli ISA

**9. Valutazione degli errori e proiezione sull'universo**

Nel caso in cui, a seguito delle verifiche effettuate sugli elementi selezionati, vengano riscontrate differenze tra saldo contabile e saldo verificato, il programma consente la determinazione automatica della stima dell'errore sull'intera popolazione oggetto di campionamento.

**9.1 Determinazione del tainting**
Per ciascun elemento con errore viene calcolato il tainting, definito come:
- Errore / Valore contabile dell'elemento

Il tainting rappresenta la percentuale di errore rispetto al valore originario dell'item selezionato.

Gli elementi privi di differenze assumono tainting pari a zero.

**9.2 Most Likely Error (MLE)**
Nel caso di utilizzo del metodo MUS, l'errore proiettato sull'universo (Most Likely Error - MLE) e determinato secondo la seguente logica:
- Intervallo di campionamento = Universo residuo / Numero di selezioni
- MLE = Intervallo x Somma dei tainting riscontrati

In forma equivalente:
- MLE = (Somma dei tainting / Numero di selezioni) x Universo residuo

Il MLE rappresenta la stima puntuale dell'errore presente nella popolazione oggetto di campionamento.

**9.3 Valore calcolato secondo il metodo selezionato**
Il programma puo inoltre calcolare un valore ulteriore derivante dal metodo di campionamento prescelto, determinato sulla base dei parametri impostati dall'utente (universo, numerosita del campione, livello di confidenza).

Tale valore e riportato a fini informativi e documentali, in coerenza con il metodo applicato.

**9.4 Interpretazione dei risultati**
Il Most Likely Error rappresenta la stima centrale dell'errore.

Il confronto tra i valori stimati e la materialita definita in fase di pianificazione consente al revisore di:
- valutare la significativita delle differenze riscontrate
- determinare l'eventuale necessita di estendere le verifiche
- supportare il giudizio professionale finale

**9.5 Responsabilita professionale**
La stima dell'errore sull'universo costituisce uno strumento di supporto all'analisi, ma non sostituisce:
- la valutazione qualitativa delle differenze
- l'analisi delle cause degli errori
- il giudizio professionale del revisore

Eventuali decisioni in merito all'estensione delle procedure o alla richiesta di rettifiche restano di esclusiva competenza del professionista.

**10. Avvertenze e limiti di utilizzo**
Il presente strumento:
- non sostituisce il giudizio professionale del revisore
- supporta, ma non determina autonomamente, le conclusioni di revisione
- deve essere utilizzato nel contesto di una corretta valutazione del rischio
- presuppone la correttezza dei dati caricati dall'utente

L'utente resta responsabile:
- della coerenza dei parametri inseriti
- della valutazione finale dei risultati
- dell'adeguatezza del campione selezionato

Il software rappresenta uno strumento di supporto operativo, conforme ai principi di revisione, ma non sostitutivo dell'attivita professionale.
"""
    )

uploaded_file = st.file_uploader("Carica file Excel o CSV", type=["xlsx", "csv"])
if uploaded_file is None:
    st.stop()

ext = os.path.splitext(uploaded_file.name)[1].lower()
if ext not in [".xlsx", ".csv"]:
    st.error("Formato file non supportato.")
    st.stop()

file_bytes = uploaded_file.getvalue()
sheet_name = None
if ext == ".xlsx":
    sheet_names = get_excel_sheet_names(file_bytes)
    if not sheet_names:
        st.error("File Excel senza fogli leggibili.")
        st.stop()
    sheet_name = st.selectbox("Foglio Excel da caricare", sheet_names, index=0)

colonne = get_file_columns(file_bytes, ext, sheet_name)
if not colonne:
    if ext == ".csv":
        st.error("CSV non valido. Usa separatore virgola, punto e virgola o tab.")
    else:
        st.error("File Excel vuoto o non leggibile.")
    st.stop()

st.sidebar.header("Informazioni revisione")
societa = st.sidebar.text_input("Societa")
revisione_al_str = st.sidebar.text_input("Revisione al (GG/MM/AAAA)", value=datetime.now().strftime("%d/%m/%Y"))
try:
    datetime.strptime(revisione_al_str, "%d/%m/%Y")
except ValueError:
    st.sidebar.warning("Formato data non valido. Uso data odierna.")
    revisione_al_str = datetime.now().strftime("%d/%m/%Y")
preparato_da = st.sidebar.text_input("Preparato da")
data_ora = datetime.now().strftime("%d/%m/%Y %H:%M")

st.sidebar.header("Mappatura colonne")
if len(colonne) < 3:
    st.error("Il file deve contenere almeno 3 colonne.")
    st.stop()
codice_col = st.sidebar.selectbox("Colonna Codice", colonne, index=0)
descr_col = st.sidebar.selectbox("Colonna Descrizione", colonne, index=1)
valore_col = st.sidebar.selectbox("Colonna Valore", colonne, index=2)
if len({codice_col, descr_col, valore_col}) < 3:
    st.warning("Le colonne devono essere diverse.")
    st.stop()

with st.spinner("Caricamento dati selezionati in corso..."):
    df = load_selected_columns(file_bytes, ext, (codice_col, descr_col, valore_col), sheet_name)

if df.empty:
    st.error("File caricato vuoto o senza righe dati.")
    st.stop()

if "mostra_grafico_universo" not in st.session_state:
    st.session_state["mostra_grafico_universo"] = False
if st.sidebar.button("Mostra/Nascondi grafico universo"):
    st.session_state["mostra_grafico_universo"] = not st.session_state["mostra_grafico_universo"]

st.sidebar.header("Parametri di campionamento")
soglia_tipo = st.sidebar.radio("Tipo soglia Key Items", ["Percentuale sul totale", "Soglia numerica", "Nessun Key Item"])
if soglia_tipo == "Percentuale sul totale":
    perc_key = st.sidebar.number_input("Soglia Key Items (%)", 0, 100, 30)
    soglia_key_num = None
elif soglia_tipo == "Soglia numerica":
    soglia_key_num = st.sidebar.number_input("Soglia Key Items (EUR)", min_value=0.0, value=10000.0, step=100.0, format="%.2f")
    perc_key = None
else:
    perc_key = None
    soglia_key_num = None

materialita = st.sidebar.number_input("Materialita (EUR)", min_value=1.0, value=1_000_000.0)
confidence_level = st.sidebar.number_input(
    "Confidence Level (%)",
    min_value=1.0,
    max_value=100.0,
    value=80.0,
    step=0.1,
    format="%.1f",
)

st.sidebar.header("Metodo selezione Items")
metodo = st.sidebar.radio("", ["MUS", "Intervallo", "Casuale"])

manual_starting_point = None
starting_point_mode = None
if metodo in ["MUS", "Intervallo"]:
    starting_point_mode = st.sidebar.radio("Selezione Starting Point", ["Automatica", "Manuale"])
    if starting_point_mode == "Manuale":
        min_starting_point = 1.0 if metodo == "Intervallo" else 0.0
        default_starting_point = 1.0 if metodo == "Intervallo" else 0.0
        manual_starting_point = st.sidebar.number_input(
            "Inserisci valore Starting Point", min_value=min_starting_point, value=default_starting_point, step=1.0, format="%.2f"
        )

errore_atteso_perc = st.sidebar.slider("Errore atteso (%)", min_value=0, max_value=70, value=0, step=1)
errore_atteso_valore = materialita * errore_atteso_perc / 100
materialita_net_benchmark = materialita - errore_atteso_valore
st.sidebar.info(
    f"Errore atteso: {format_number_it(errore_atteso_valore, 0)} EUR\n\n"
    f"Materialita netta per campionamento: {format_number_it(materialita_net_benchmark, 0)} EUR"
)
calcola_campione = st.sidebar.button("Calcola selezione campione")

df[valore_col] = df[valore_col].apply(parse_number_it)
valori_non_validi = int(df[valore_col].isna().sum())
df = df.dropna(subset=[valore_col]).reset_index(drop=True)
if valori_non_validi > 0:
    st.warning(f"Scartate {valori_non_validi} righe con valore non numerico.")

df_base = df.copy()
df_mus = df.sort_values(by=valore_col, ascending=False).reset_index(drop=True)

tot_items = len(df_base)
tot_valore = float(df_base[valore_col].sum())
top5_val = float(df_base.sort_values(by=valore_col, ascending=False).head(5)[valore_col].sum())
top5_perc = top5_val / tot_valore * 100 if tot_valore != 0 and tot_items >= 5 else 100.0

st.subheader("Universo completo")
preview_rows = 1000
df_preview = df_base.head(preview_rows)
st.dataframe(
    format_display_df(df_preview, {valore_col: 2}),
    use_container_width=True,
)
if len(df_base) > preview_rows:
    st.caption(f"Visualizzate le prime {preview_rows} righe su {len(df_base):,} totali.")
c1, c2, c3 = st.columns(3)
c1.metric("Totale items", tot_items)
c2.metric("Valore totale", format_number_it(tot_valore, 0))
c3.metric("Top 5 (% valore)", f"{top5_perc:.2f}%")

if st.session_state.get("mostra_grafico_universo"):
    render_universe_stratification(df_base, valore_col)

selection_signature = (
    uploaded_file.name,
    len(file_bytes),
    sheet_name,
    codice_col,
    descr_col,
    valore_col,
    soglia_tipo,
    perc_key,
    soglia_key_num,
    materialita,
    confidence_level,
    metodo,
    starting_point_mode,
    manual_starting_point,
    errore_atteso_perc,
)
if st.session_state.get("selection_signature") != selection_signature:
    for key in [
        "key_items",
        "items_selezionati",
        "num_items_teorici",
        "starting_point",
        "intervallo_utilizzato",
        "errori_editor",
    ]:
        st.session_state.pop(key, None)
    st.session_state["selection_signature"] = selection_signature

if calcola_campione or ("key_items" in st.session_state and "items_selezionati" in st.session_state):
    if calcola_campione:
        if "errori_editor" in st.session_state:
            del st.session_state["errori_editor"]

        df_sorted = df_mus.copy()
        if soglia_tipo == "Percentuale sul totale":
            soglia_key = tot_valore * perc_key / 100
            df_sorted["cumulativo"] = df_sorted[valore_col].cumsum()
            key_items = df_sorted[df_sorted["cumulativo"] <= soglia_key].copy()
        elif soglia_tipo == "Soglia numerica":
            key_items = df_sorted[df_sorted[valore_col] > soglia_key_num].copy()
        else:
            key_items = df_sorted.iloc[0:0].copy()

        residuo = df_sorted.drop(key_items.index).copy()
        residuo_tot = float(residuo[valore_col].sum())
        confidence_factor = 100 * (1 - ((100 - confidence_level) / 100) ** (1 / 100))
        num_items = int(round((residuo_tot / materialita_net_benchmark) * confidence_factor))
        num_items = max(num_items, 1)
        num_items_teorici = num_items

        selected_items = pd.DataFrame(columns=df_sorted.columns)
        starting_point = None

        if metodo == "MUS" and residuo_tot > 0:
            np.random.seed(42)
            residuo = df_mus.drop(key_items.index)
            intervallo = residuo_tot / num_items
            if starting_point_mode == "Manuale" and manual_starting_point is not None:
                starting_point = manual_starting_point
            else:
                starting_point = np.random.uniform(0, intervallo)

            residuo["cumulativo"] = residuo[valore_col].cumsum()
            soglie = [starting_point + i * intervallo for i in range(num_items)]
            idx = set()
            for s in soglie:
                r = residuo[residuo["cumulativo"] >= s].head(1)
                if not r.empty:
                    idx.add(r.index[0])
            selected_items = residuo.loc[list(idx)]
        elif metodo == "Intervallo":
            residuo = df_base.drop(key_items.index)
            step = max(1, len(residuo) // num_items)
            if starting_point_mode == "Manuale" and manual_starting_point is not None:
                start_idx = max(0, int(round(manual_starting_point)) - 1)
                selected_items = residuo.iloc[start_idx::step].head(num_items)
                starting_point = start_idx + 1
            else:
                np.random.seed(42)
                start_idx = np.random.randint(0, step)
                selected_items = residuo.iloc[start_idx::step].head(num_items)
                starting_point = start_idx + 1
        else:
            residuo = df_base.drop(key_items.index)
            np.random.seed(42)
            selected_items = residuo.sample(n=min(num_items, len(residuo)), random_state=42)

        if metodo == "MUS":
            intervallo_utilizzato = residuo_tot / num_items_teorici if num_items_teorici > 0 else 0.0
        else:
            intervallo_utilizzato = residuo_tot / len(selected_items) if len(selected_items) > 0 else 0.0
        st.session_state["key_items"] = key_items
        st.session_state["items_selezionati"] = selected_items
        st.session_state["num_items_teorici"] = num_items_teorici
        st.session_state["starting_point"] = starting_point
        st.session_state["intervallo_utilizzato"] = intervallo_utilizzato
    else:
        key_items = st.session_state["key_items"]
        selected_items = st.session_state["items_selezionati"]
        num_items_teorici = st.session_state.get("num_items_teorici", len(selected_items))
        starting_point = st.session_state.get("starting_point")
        intervallo_utilizzato = st.session_state.get("intervallo_utilizzato", 0.0)

    riepilogo = pd.DataFrame(
        {
            "Categoria": ["Universo", "Key Items", "Selezioni teoriche", "Items unici selezionati"],
            "Numero items": [tot_items, len(key_items), num_items_teorici, len(selected_items)],
            "Valore (EUR)": [
                euro(tot_valore),
                euro(key_items[valore_col].sum()),
                "-",
                euro(selected_items[valore_col].sum()),
            ],
            "% su totale": [
                "100.00%",
                f"{key_items[valore_col].sum()/tot_valore*100:.2f}%" if tot_valore else "0.00%",
                "-",
                f"{selected_items[valore_col].sum()/tot_valore*100:.2f}%" if tot_valore else "0.00%",
            ],
            "Starting point": ["-", "-", "-", f"{starting_point:.2f}" if starting_point is not None else "-"],
        }
    )
    st.subheader("Riepilogo selezione")
    st.table(riepilogo)
    selezioni_duplicate_mus = max(0, num_items_teorici - len(selected_items)) if metodo == "MUS" else 0
    if metodo == "MUS" and num_items_teorici > len(selected_items):
        st.info(
            f"Nel metodo MUS {num_items_teorici} selezioni teoriche hanno prodotto "
            f"{len(selected_items)} item unici: {selezioni_duplicate_mus} selezioni "
            "sono ricadute su item gia selezionati. "
            "Il revisore puo valutare, sulla base del proprio giudizio professionale, "
            "se integrare il campione con ulteriori item dal residuo non selezionato, "
            "fino a raggiungere il numero teorico di selezioni."
        )

    st.subheader("Key Items")
    st.dataframe(
        format_display_df(key_items[[codice_col, descr_col, valore_col]], {valore_col: 2}),
        use_container_width=True,
    )
    st.subheader("Items selezionati")
    st.dataframe(
        format_display_df(selected_items[[codice_col, descr_col, valore_col]], {valore_col: 2}),
        use_container_width=True,
    )
    if soglia_tipo == "Nessun Key Item" and metodo == "MUS" and not selected_items.empty:
        max_valore = float(df_base[valore_col].max())
        if (selected_items[valore_col] == max_valore).any():
            st.info(
                "L'item di importo maggiore non e stato classificato come Key Item: "
                "e stato selezionato dal metodo MUS come elemento del campione, comportamento fisiologico per questo metodo."
            )

    st.subheader("Errori emersi e proiezione sull'universo")
    saldo_col = "Saldo corretto emerso (EUR)"
    errore_col = "Errore (EUR)"
    errore_perc_col = "Errore (%)"

    errori_df = selected_items[[codice_col, descr_col, valore_col]].copy()
    errori_df[saldo_col] = errori_df[valore_col]
    errori_editati = st.data_editor(
        errori_df,
        key="errori_editor",
        hide_index=True,
        use_container_width=True,
        column_config={
            valore_col: st.column_config.NumberColumn(label=valore_col, format="%.2f", disabled=True),
            saldo_col: st.column_config.NumberColumn(label=saldo_col, format="%.2f"),
        },
    )

    errori_editati[saldo_col] = pd.to_numeric(errori_editati[saldo_col], errors="coerce").fillna(0.0)
    valori_originali = pd.to_numeric(errori_editati[valore_col], errors="coerce").fillna(0.0)
    saldi_corretti = pd.to_numeric(errori_editati[saldo_col], errors="coerce").fillna(0.0)
    errore_importo = saldi_corretti - valori_originali
    errore_perc = np.where(valori_originali != 0, (errore_importo / valori_originali) * 100, 0.0)
    somma_perc_errori = float(np.nansum(errore_perc))
    somma_tainting_rate = somma_perc_errori / 100

    dettaglio_errori = errori_editati[[codice_col, descr_col, valore_col, saldo_col]].copy()
    dettaglio_errori[errore_col] = errore_importo
    dettaglio_errori[errore_perc_col] = errore_perc
    dettaglio_errori_con_errori = dettaglio_errori[np.abs(dettaglio_errori[errore_col]) > 1e-12]
    dettaglio_errori_export = dettaglio_errori_con_errori.copy()
    if not dettaglio_errori_export.empty:
        totale_valore_errori = float(dettaglio_errori_export[valore_col].sum())
        totale_errore_errori = float(dettaglio_errori_export[errore_col].sum())
        totale_errori_row = {
            codice_col: "Totale",
            descr_col: "",
            valore_col: totale_valore_errori,
            saldo_col: dettaglio_errori_export[saldo_col].sum(),
            errore_col: totale_errore_errori,
            errore_perc_col: np.nan,
        }
        dettaglio_errori_export = pd.concat(
            [dettaglio_errori_export, pd.DataFrame([totale_errori_row])],
            ignore_index=True,
        )
    universo_no_key_items = max(0.0, float(tot_valore - key_items[valore_col].sum()))
    valore_campione_selezionato = float(selected_items[valore_col].sum())

    if metodo == "MUS":
        mle = somma_tainting_rate * intervallo_utilizzato
        formula_mle = (
            "Calcolo MLE (MUS): Somma tainting rate x Intervallo utilizzato = "
            f"{somma_tainting_rate:.6f} x {format_number_it(intervallo_utilizzato, 2)} = {format_number_it(mle, 2)}"
        )
    else:
        totale_errore = float(np.nansum(errore_importo))
        tasso_errore = totale_errore / valore_campione_selezionato if valore_campione_selezionato else 0.0
        mle = tasso_errore * universo_no_key_items
        formula_mle = (
            "Calcolo MLE (metodo non MUS): Errore totale / Totale campione selezionato x Universo (no key items) = "
            f"{format_number_it(totale_errore, 2)} / {format_number_it(valore_campione_selezionato, 2)} x "
            f"{format_number_it(universo_no_key_items, 2)} = {format_number_it(mle, 2)}"
        )

    if not dettaglio_errori_con_errori.empty:
        st.write("Dettaglio items con errore")
        st.dataframe(
            format_display_df(
                dettaglio_errori_export,
                {valore_col: 2, saldo_col: 2, errore_col: 2, errore_perc_col: 4},
            ),
            use_container_width=True,
        )
    else:
        st.write("Nessun errore rilevato negli items selezionati.")

    m1, m2, m3, m4, m5 = st.columns(5)
    m1.metric("Universo (no key items)", format_number_it(universo_no_key_items, 0))
    m2.metric("Items effettivi selezionati", len(errori_editati))
    m3.metric("Valore items selezionati", format_number_it(valore_campione_selezionato, 0))
    m4.metric("Intervallo utilizzato", format_number_it(intervallo_utilizzato, 0))
    m5.metric("Most likely error (MLE)", format_number_it(mle, 0))
    st.info(formula_mle)
    st.write("nella determinazione del MLE si sono considerati sia gli errori negativi che gli errori positivi")

    excel_buffer = BytesIO()
    with pd.ExcelWriter(excel_buffer, engine="xlsxwriter") as writer:
        info_df = pd.DataFrame(
            {
                "Parametro": [
                    "Societa",
                    "Revisione al",
                    "Preparato da",
                    "Materialita",
                    "Errore atteso (%)",
                    "Errore atteso (EUR)",
                    "Materialita netta per campionamento",
                    "Confidence level",
                    "Metodo selezione Items",
                    "Starting point",
                ],
                "Valore": [
                    societa,
                    revisione_al_str,
                    preparato_da,
                    euro(materialita),
                    f"{errore_atteso_perc}%",
                    euro(errore_atteso_valore),
                    euro(materialita_net_benchmark),
                    f"{confidence_level}%",
                    metodo,
                    f"{starting_point:.2f}" if starting_point is not None else "-",
                ],
            }
        )
        info_df.to_excel(writer, sheet_name="Riepilogo", index=False, startrow=0)
        riepilogo.to_excel(writer, sheet_name="Riepilogo", index=False, startrow=info_df.shape[0] + 2)
        key_items[[codice_col, descr_col, valore_col]].to_excel(writer, sheet_name="Key Items", index=False)
        selected_items[[codice_col, descr_col, valore_col]].to_excel(writer, sheet_name="Items selezionati", index=False)

        mle_info_rows = [
            ("Universo (no key items)", format_number_it(universo_no_key_items, 0)),
            ("Items effettivi selezionati", len(errori_editati)),
            ("Valore items selezionati", valore_campione_selezionato),
            ("Intervallo utilizzato", intervallo_utilizzato),
        ]
        if metodo == "MUS":
            mle_info_rows.append(("Somma tainting rate", somma_tainting_rate))
        mle_info_rows.extend(
            [
                ("Calcolo MLE", formula_mle),
                ("Most likely error (MLE)", mle),
                ("Nota", "nella determinazione del MLE si sono considerati sia gli errori negativi che gli errori positivi"),
            ]
        )
        mle_info_df = pd.DataFrame(mle_info_rows, columns=["Parametro", "Valore"])
        mle_info_df.to_excel(writer, sheet_name="Errori e MLE", index=False, startrow=0)
        dettaglio_errori_export.to_excel(writer, sheet_name="Errori e MLE", index=False, startrow=mle_info_df.shape[0] + 2)
    excel_buffer.seek(0)

    st.download_button(
        "Export Excel",
        data=excel_buffer,
        file_name="audit_sampling.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

    def export_word():
        doc = Document()
        doc.add_heading("MEMO DI REVISIONE - SELEZIONE CAMPIONE", 0)
        doc.add_paragraph(f"Societa: {societa}")
        doc.add_paragraph(f"Revisione al: {revisione_al_str}")
        doc.add_paragraph(f"File: {uploaded_file.name}")
        doc.add_paragraph(f"Data: {data_ora}")
        doc.add_paragraph(f"Preparato da: {preparato_da}")
        doc.add_paragraph(f"Metodo: {metodo}")
        if starting_point is not None:
            doc.add_paragraph(f"Starting point: {starting_point:.2f}")
        doc.add_paragraph(f"Materialita: {euro(materialita)}  |  Confidence level: {confidence_level}%")
        doc.add_paragraph(
            f"Errore atteso: {errore_atteso_perc}% ({euro(errore_atteso_valore)})  |  "
            f"Materialita netta per campionamento: {euro(materialita_net_benchmark)}"
        )
        doc.add_paragraph("\nRiepilogo selezione")
        t = doc.add_table(rows=riepilogo.shape[0] + 1, cols=riepilogo.shape[1])
        for j, col in enumerate(riepilogo.columns):
            t.cell(0, j).text = str(col)
        for i in range(riepilogo.shape[0]):
            for j in range(riepilogo.shape[1]):
                t.cell(i + 1, j).text = str(riepilogo.iloc[i, j])

        doc.add_paragraph("\nMost likely error (MLE)")
        doc.add_paragraph(f"Universo (no key items): {format_number_it(universo_no_key_items, 0)}")
        doc.add_paragraph(f"Items effettivi selezionati: {len(errori_editati)}")
        doc.add_paragraph(f"Valore items selezionati: {format_number_it(valore_campione_selezionato, 0)}")
        doc.add_paragraph(f"Intervallo utilizzato: {format_number_it(intervallo_utilizzato, 0)}")
        if metodo == "MUS":
            doc.add_paragraph(f"Somma tainting rate: {somma_tainting_rate:.6f}")
        doc.add_paragraph(formula_mle)
        doc.add_paragraph(f"Most likely error (MLE): {format_number_it(mle, 0)}")
        doc.add_paragraph("nella determinazione del MLE si sono considerati sia gli errori negativi che gli errori positivi")

        doc.add_paragraph("\nDettaglio items con errore")
        headers = [codice_col, descr_col, valore_col, saldo_col, errore_col, errore_perc_col]
        te = doc.add_table(rows=len(dettaglio_errori_export) + 1, cols=len(headers))
        for j, h in enumerate(headers):
            te.cell(0, j).text = h
        for i, row in enumerate(dettaglio_errori_export[headers].itertuples(index=False)):
            for j, v in enumerate(row):
                col_name = headers[j]
                if col_name in [valore_col, saldo_col, errore_col, errore_perc_col]:
                    num_v = pd.to_numeric(v, errors="coerce")
                    te.cell(i + 1, j).text = "" if pd.isna(num_v) else f"{num_v:.2f}"
                else:
                    te.cell(i + 1, j).text = str(v)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    st.download_button(
        "Export Word",
        data=export_word(),
        file_name="audit_sampling.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
